import os
import fitz  # PyMuPDF
import docx  # python-docx
import requests
import chromadb
import concurrent.futures
import re
import pytesseract
from PIL import Image
import pandas as pd
import uuid

OLLAMA_URL = "http://localhost:11434/api/embeddings"
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DB_PATH = os.path.join(BASE_DIR, "chroma_db")

chroma_client = chromadb.PersistentClient(path=DB_PATH)
collection = chroma_client.get_or_create_collection(name="brd_docs_nomic")

# Windows ke liye Tesseract OCR ka exact path (Scanned images aur PDFs se text read karne ke liye)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def get_embedding(text):
    try:
        # Using a dedicated, fast embedding model
        response = requests.post(OLLAMA_URL, json={"model": "nomic-embed-text", "prompt": text})
        if response.status_code == 200:
            return response.json().get("embedding", [])
        else:
            print("Ollama Embedding Error:", response.text)
    except Exception as e:
        print("Embedding Error:", e)
    return []

def process_document(file_path, user_id):
    ext = os.path.splitext(file_path)[1].lower()
    raw_text = ""
    
    if ext == ".pdf":
        def process_pdf_page(page_num):
            local_doc = fitz.open(file_path)
            page = local_doc[page_num]
            page_text = page.get_text("text")
            
            # Agar page mein text nahi mila (mtlb image hai), toh OCR perform karein
            if len(page_text.strip()) < 50:
                try:
                    # Lower DPI for faster OCR (e.g., 100 instead of 150)
                    pix = page.get_pixmap(dpi=100)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    page_text += "\n" + pytesseract.image_to_string(img)
                except Exception as e:
                    print(f"OCR Error on page {page_num}: {e}")
            
            raw_text_part = page_text + "\n"
            # Extract tables into Markdown format
            if hasattr(page, "find_tables"):
                for table in page.find_tables():
                    extracted = table.extract()
                    for i, row in enumerate(extracted):
                        clean_row = [str(cell).replace('\n', ' ').strip() if cell is not None else "" for cell in row]
                        raw_text_part += "| " + " | ".join(clean_row) + " |\n"
                        if i == 0:
                            raw_text_part += "|" + "|".join(["---"] * len(clean_row)) + "|\n"
                    raw_text_part += "\n"
            local_doc.close()
            return raw_text_part

        # PyMuPDF Documents aren't thread-safe, so we open a local copy in each thread
        doc = fitz.open(file_path)
        num_pages = doc.page_count
        doc.close()

        # Multi-threading for parallel page extraction
        with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
            # map() preserves the correct page order
            results = executor.map(process_pdf_page, range(num_pages))
            raw_text = "".join(results)
    elif ext == ".docx":
        doc = docx.Document(file_path)
        raw_text = "\n".join([para.text for para in doc.paragraphs]) + "\n\n"
        # Extract tables into Markdown format
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                # Get cell text and remove newlines so the markdown table doesn't break
                clean_row = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                raw_text += "| " + " | ".join(clean_row) + " |\n"
                if i == 0: # Add table header separator
                    raw_text += "|" + "|".join(["---"] * len(clean_row)) + "|\n"
            raw_text += "\n"
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            raw_text = f.read()
    elif ext == ".xlsx":
        try:
            xls = pd.ExcelFile(file_path)
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                df = df.fillna("") # Khali cells ko blank banayein
                headers = [str(c).replace('\n', ' ').strip() for c in df.columns]
                
                # Chunking logic: 20 rows per chunk with repeating headers
                batch_size = 20
                for i in range(0, len(df), batch_size):
                    batch_df = df.iloc[i:i+batch_size]
                    raw_text += f"### Sheet: {sheet_name} (Rows {i+1} to {i+len(batch_df)})\n"
                    raw_text += "| " + " | ".join(headers) + " |\n"
                    raw_text += "|" + "|".join(["---"] * len(headers)) + "|\n"
                    for _, row in batch_df.iterrows():
                        row_vals = [str(val).replace('\n', ' ').strip() for val in row]
                        raw_text += "| " + " | ".join(row_vals) + " |\n"
                    raw_text += "\n\n"
        except Exception as e:
            print(f"Excel Error: {e}")
    else:
        return 0 # Unsupported format
        
    # 1. Basic Cleaning: Extra spaces aur tabs ko single space banayein
    cleaned_text = re.sub(r'[ \t]+', ' ', raw_text)
    # 2. Multiple newlines ko paragraph break mein badlein
    cleaned_text = re.sub(r'\n\s*\n', '\n\n', cleaned_text)
    
    # 3. Noise Removal: Chhoti garbage lines hatayein par Markdown tables bachayein
    lines = cleaned_text.split('\n')
    good_lines = []
    for line in lines:
        s_line = line.strip()
        if (s_line.startswith('|') and s_line.endswith('|')) or len(s_line) > 3 or s_line.isalpha():
            good_lines.append(s_line)
            
    text = "\n".join(good_lines) + "\n\n"

    # Advanced Chunking Strategy with Overlap
    chunk_size = 2500  # Increased chunk size to reduce the total number of chunks (makes processing much faster)
    chunk_overlap = 300
    step_size = chunk_size - chunk_overlap
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), step_size)]

    def process_chunk(index, chunk_text):
        if len(chunk_text.strip()) < 10: return None
        emb = get_embedding(chunk_text)
        if emb:
            return {
                "id": f"{user_id}_{os.path.basename(file_path)}_{index}",
                "embedding": emb,
                "document": chunk_text,
                "metadata": {"user_id": user_id, "source": os.path.basename(file_path)}
            }
        return None

    # Multithreading se processing fast karenge (Max workers increased for faster Ollama API calls)
    with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
        results = executor.map(lambda p: process_chunk(p[0], p[1]), enumerate(chunks))
        valid_data = [res for res in results if res is not None]

    # Ek hi baar mein (Batch Upsert) saare chunks DB mein save karenge
    if valid_data:
        collection.upsert(
            ids=[item["id"] for item in valid_data],
            embeddings=[item["embedding"] for item in valid_data],
            documents=[item["document"] for item in valid_data],
            metadatas=[item["metadata"] for item in valid_data]
        )
        
    return len(valid_data), text

def save_chat_memory(user_id, query, response):
    # Create a structured memory block
    text = f"Past Conversation -> User: {query} \nAI: {response}"
    if len(text) < 30: return # Faltu chhote messages (jaise "hi", "ok") save na karein
    text = text[:1500] # Limit size for fast embedding
    emb = get_embedding(text)
    if emb:
        doc_id = f"mem_{user_id}_{uuid.uuid4().hex[:10]}"
        try:
            collection.upsert(
                ids=[doc_id],
                embeddings=[emb],
                documents=[text],
                metadatas=[{"user_id": user_id, "source": "Past Chat Memory"}]
            )
        except Exception as e:
            print(f"Memory Save Error: {e}")

def retrieve_context(query, user_id, top_k=3):
    query_emb = get_embedding(query)
    if not query_emb: return "", []
    
    try:
        # Filter out old chat memory to keep new sessions clean and lightning fast
        results = collection.query(
            query_embeddings=[query_emb], 
            n_results=top_k, 
            where={"$and": [{"user_id": user_id}, {"source": {"$ne": "Past Chat Memory"}}]}
        )
        if results and results['documents'] and results['documents'][0]:
            context = "\n...\n".join(results['documents'][0])
            sources = list(set([meta["source"] for meta in results['metadatas'][0] if "source" in meta]))
            return context, sources
    except: pass
    return "", []